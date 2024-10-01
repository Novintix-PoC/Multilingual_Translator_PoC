import streamlit as st
from docx import Document
from docx.shared import Pt
from transformers import M2M100ForConditionalGeneration, M2M100Tokenizer
import datetime
import os
from pdf2docx import parse
import docx2pdf
import tempfile
from io import BytesIO

st.markdown("""
    <style>
        @import url('https://fonts.googleapis.com/css2?family=Roboto:wght@300;400;700&display=swap');
    
        .stApp {
            background: linear-gradient(to top, #0e4166 40%, #000000 90%); #for background;
            font-family: 'Roboto', sans-serif;
        }

        [data-testid=stSidebar] {
            background: linear-gradient(to top, #0e4166 60%, #000000 90%); # for sidebar;
            margin-top:60px;
        }

        [data-testid=stHeader] {
            background: linear-gradient( #0e4166 20%, #000000 100%); #for header;
        }
        [data-testid=stWidgetLabel] {
            color:#f4a303; # for small headings;
        }

        .stHeading h1{ 
            color:#f4a303; # for Heading "select Languages";
        }

        [data-testid=stAppViewBlockContainer],[data-testid=stVerticalBlock]{
            margin-top:30px; # for centering the content and sidebar;
        }

        .title-container {
            background: rgb(0,27,44);
            backdrop-filter: blur(10px);
            border-radius: 30px;
            text-align:center;
            padding:5px;
            margin-bottom:10px;
            box-shadow: 0px 4px 6px rgba(0, 0, 0, 0.1);
        }
   
    </style>""", unsafe_allow_html=True)

header_html = f"""
    <div class="title-container">
        <span style="font-size: 30px; font-weight: bold;color:#023d59;">Welcome to</span>
        <img src="https://novintix.com/wp-content/uploads/2023/12/logov2.png" style="height: 45px;padding:5px;  margin-bottom: 10px;">
        <span style="font-size: 30px; font-weight: bold; color:#f4a303;">Multilingual Language Translation Tool</span>
    </div>
    """

st.markdown(header_html, unsafe_allow_html=True)

# Mapping of full language names to language codes
language_mapping = {
    "Bulgarian": "bg", "Chinese": "zh", "Croatian": "hr", "Czech": "cs", "Danish": "da",
    "Dutch": "nl", "Estonian": "et", "English": "en", "Finnish": "fi", "French": "fr",
    "German": "de", "Greek": "el", "Hungarian": "hu", "Icelandic": "is", "Indonesian": "id",
    "Italian": "it", "Kazakh": "kk", "Korean": "ko", "Latvian": "lv", "Lithuanian": "lt",
    "Macedonian": "mk", "Norwegian": "no", "Polish": "pl", "Portuguese": "pt", "Romanian": "ro",
    "Russian": "ru", "Serbian": "sr", "Slovak": "sk", "Slovenian": "sl", "Spanish": "es",
    "Swedish": "sv", "Turkish": "tr", "Vietnamese": "vi"
}

def load_translation_model():
    model_name = 'facebook/m2m100_418M'
    tokenizer = M2M100Tokenizer.from_pretrained(model_name)
    model = M2M100ForConditionalGeneration.from_pretrained(model_name)
    return tokenizer, model

def translate_text(text: str, src_lang: str, tgt_lang: str, tokenizer, model):
    tokenizer.src_lang = src_lang
    encoded = tokenizer(text, return_tensors="pt", padding=True, truncation=True)
    generated_tokens = model.generate(**encoded, forced_bos_token_id=tokenizer.get_lang_id(tgt_lang))
    translated_text = tokenizer.batch_decode(generated_tokens, skip_special_tokens=True)[0]
    return translated_text

def copy_run_format(source_run, target_run):
    target_run.bold = source_run.bold
    target_run.italic = source_run.italic
    target_run.underline = source_run.underline
    target_run.font.name = source_run.font.name
    target_run.font.size = source_run.font.size
    target_run.font.color.rgb = source_run.font.color.rgb
    text = source_run.text
    if text.isupper():
        target_run.text = text.upper()
    elif text.islower():
        target_run.text = text.lower()

def translate_text_with_format(paragraph, src_lang, tgt_lang, tokenizer, model):
    new_runs = []
    for run in paragraph.runs:
        translated_text = translate_text(run.text, src_lang, tgt_lang, tokenizer, model)
        if "™" in run.text:
            translated_text = translated_text.replace("TM", "™")
        new_run = paragraph.add_run(translated_text)
        copy_run_format(run, new_run)
        run.clear()

def convert_pdf_to_docx(pdf_path, docx_path):
    parse(pdf_path, docx_path, start=0, end=None)
    print("PDF to DOCX Done...")

def translate_docx(doc_path, src_lang, tgt_langs, temp_dir, input_file_type):
    start_time = datetime.datetime.now()
    tokenizer, model = load_translation_model()
    translated_files = {}

    for tgt_lang in tgt_langs:
        doc = Document(doc_path)

        for para in doc.paragraphs:
            if para.text.strip():
                translate_text_with_format(para, language_mapping[src_lang], language_mapping[tgt_lang], tokenizer, model)

        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        cell.text = translate_text(cell.text, language_mapping[src_lang], language_mapping[tgt_lang], tokenizer, model)

        output_docx_path = os.path.join(temp_dir, f"translated_{src_lang}_to_{tgt_lang}.docx")
        doc.save(output_docx_path)
        os.remove(output_docx_path)
        translated_files[tgt_lang] = output_docx_path

    end_time = datetime.datetime.now()
    time_diff = end_time - start_time
    hours, remainder = divmod(time_diff.seconds, 3600)
    minutes, seconds = divmod(remainder, 60)

    print(f"Translation completed in {hours} hours, {minutes} minutes, and {seconds} seconds.")

    return translated_files

def main():
    st.sidebar.title("Select Languages")
    src_lang = st.sidebar.selectbox("Select source language:", list(language_mapping.keys()))
    tgt_langs = st.sidebar.multiselect("Select target languages:", list(language_mapping.keys()))

    input_file = st.file_uploader("Upload Word or PDF file for translation:", type=["docx", "pdf"])

    if st.button("Translate"):
        if not input_file:
            st.warning("Please upload a Word or PDF file for translation.")
        elif not src_lang:
            st.warning("Please select a source language.")
        elif not tgt_langs:
            st.warning("Please select at least one target language.")
        else:
            with st.spinner("Translating..."):
                input_file_type = input_file.name.split('.')[-1].lower()
                
                with tempfile.TemporaryDirectory() as temp_dir:
                    temp_input_path = os.path.join(temp_dir, input_file.name)
                    with open(temp_input_path, 'wb') as f:
                        f.write(input_file.getvalue())
                    
                    temp_docx_path = os.path.join(temp_dir, "temp.docx")
                    if input_file_type == 'pdf':
                        convert_pdf_to_docx(temp_input_path, temp_docx_path)
                    else:
                        temp_docx_path = temp_input_path

                    translated_files = translate_docx(temp_docx_path, src_lang, tgt_langs, temp_dir, input_file_type)
                    
                    st.success("Translation completed. You can now download the files.")
                    
                    for lang, file_path in translated_files.items():
                        with open(file_path, 'rb') as file:
                            file_content = file.read()
                        st.download_button(
                            label=f"Download {lang} translation",
                            data=file_content,
                            file_name=os.path.basename(file_path),
                            mime='application/octet-stream'
                        )

if __name__ == "__main__":
    main()
